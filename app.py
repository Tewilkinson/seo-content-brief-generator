import streamlit as st
import pandas as pd
import numpy as np
import requests
from bs4 import BeautifulSoup
from openai import OpenAI
from docx import Document
from io import BytesIO
import time

# ------------------------------
# Configuration
# ------------------------------

st.set_page_config(page_title="SEO Content Brief Generator", layout="wide")
st.title("SEO Content Brief Generator")

# Initialize OpenAI client using Streamlit secrets
openai_client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

# ------------------------------
# Helper Functions
# ------------------------------

def fetch_serp_top_pages(keyword, num_results=2):
    """Fetch top ranking URLs from SERPAPI for a given keyword."""
    api_key = st.secrets["SERPAPI_KEY"]
    params = {
        "engine": "google",
        "q": keyword,
        "num": num_results,
        "api_key": api_key,
        "hl": "en",
        "gl": "us"
    }
    response = requests.get("https://serpapi.com/search", params=params)
    results = response.json()
    urls = []
    try:
        for r in results.get("organic_results", []):
            urls.append(r.get("link"))
    except Exception:
        pass
    return urls

def scrape_page_h2s_and_links(url):
    """Scrape H2 sections and links from a page."""
    try:
        res = requests.get(url, timeout=10)
        soup = BeautifulSoup(res.text, "lxml")
        h2s = [h.get_text().strip() for h in soup.find_all("h2")]
        outlinks = [a['href'] for a in soup.find_all('a', href=True)]
        return h2s, outlinks
    except Exception:
        return [], []

def generate_embeddings(text_list):
    """Generate embeddings using OpenAI."""
    response = openai_client.embeddings.create(
        model="text-embedding-3-large",
        input=text_list
    )
    embeddings = [e.embedding for e in response.data]
    return embeddings

def generate_h2_content(h2, keyword):
    """Generate content for each H2 section using OpenAI."""
    prompt = f"Write a 150-200 word informational paragraph for the H2 section '{h2}' with keyword '{keyword}':"
    resp = openai_client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=300
    )
    return resp.choices[0].message.content.strip()

def generate_faqs(keyword):
    """Generate 3 FAQs and answers using OpenAI."""
    prompt = f"Generate 3 frequently asked questions with answers about '{keyword}' in short paragraphs:"
    resp = openai_client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=400
    )
    content = resp.choices[0].message.content.strip()
    faqs = []
    for q in content.split("\n"):
        if "?" in q:
            faqs.append(q.strip())
    return faqs

def create_docx(brief, keyword):
    """Generate DOCX file from the brief table."""
    doc = Document()
    doc.add_heading(f"SEO Content Brief for '{keyword}'", level=0)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Section"
    hdr_cells[1].text = "Output"
    hdr_cells[2].text = "How to fill"
    for section in brief:
        row_cells = table.add_row().cells
        row_cells[0].text = section
        row_cells[1].text = brief[section]["output"]
        row_cells[2].text = brief[section]["how"]
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ------------------------------
# User Inputs
# ------------------------------

st.sidebar.header("Settings")
keyword = st.sidebar.text_input("Primary Keyword")
uploaded_file = st.sidebar.file_uploader("Upload CSV with URLs to scan", type=["csv"])
num_top_pages = st.sidebar.slider("Number of top pages to analyze (SERP)", 1, 5, 2)

if keyword and uploaded_file:

    df_urls = pd.read_csv(uploaded_file)
    if "url" not in df_urls.columns:
        st.error("CSV must contain a column named 'url'")
    else:
        urls_to_scan = df_urls["url"].tolist()
        st.subheader("Progress")
        progress_text = st.empty()
        progress_bar = st.progress(0)

        # ------------------------------
        # Stage 1: Fetch top SERP pages
        # ------------------------------
        progress_text.text("Fetching top SERP pages...")
        time.sleep(0.5)
        serp_urls = fetch_serp_top_pages(keyword, num_top_pages)
        progress_bar.progress(10)

        # ------------------------------
        # Stage 2: Scrape pages & generate internal links
        # ------------------------------
        progress_text.text("Scraping uploaded URLs and generating semantic internal links...")
        internal_links = []
        total_h2s = []
        outlinks_count = 0
        for i, url in enumerate(urls_to_scan):
            h2s, outlinks = scrape_page_h2s_and_links(url)
            total_h2s.extend(h2s)
            internal_links.extend([l for l in outlinks if keyword.lower() in l.lower()])
            outlinks_count += len(outlinks)
            progress_bar.progress(10 + int(40 * (i+1)/len(urls_to_scan)))

        # ------------------------------
        # Stage 3: Generate H2 content
        # ------------------------------
        progress_text.text("Generating H2 section content...")
        h2_content = {}
        for i, h2 in enumerate(total_h2s):
            h2_content[h2] = generate_h2_content(h2, keyword)
            progress_bar.progress(50 + int(20 * (i+1)/len(total_h2s)))

        # ------------------------------
        # Stage 4: Generate FAQs
        # ------------------------------
        progress_text.text("Generating FAQs...")
        faqs = generate_faqs(keyword)
        progress_bar.progress(75)

        # ------------------------------
        # Stage 5: Build Brief Table
        # ------------------------------
        brief = {
            "Title": {"output": f"{keyword}: Everything You Need to Know", "how": "SEO optimized title"},
            "Meta description": {"output": f"Learn about {keyword}, how it works, and why it matters.", "how": "Unique meta description"},
            "H1": {"output": f"What is {keyword}?", "how": "Include primary keyword"},
            "Navigational side bar": {"output": "\n".join(total_h2s), "how": "Bulleted list of sections"},
            "People also asks": {"output": "\n".join(faqs), "how": "Pull from PAA or AI-generated FAQs"},
            "Top ranking pages": {"output": "\n".join(serp_urls), "how": "Top pages from SERP"},
            "URL structure": {"output": f"https://www.example.com/{keyword.replace(' ', '-').lower()}", "how": "Suggested URL structure"},
            "Est. Number of outlinks": {"output": str(outlinks_count), "how": "Number of outlinks excluding nav/footer"},
            "Internal link recommendations": {"output": "\n".join(internal_links), "how": "Semantic internal links"},
            "Number of page sections": {"output": str(len(total_h2s)+1), "how": "Based on competitor H2s, add one extra"},
        }

        # H2 Sections
        for idx, h2 in enumerate(total_h2s[:5]):  # limit to top 5 H2s
            key = f"h2_{idx}"
            brief[f"H2 section {idx+1}"] = {"output": h2_content[h2], "how": "Auto-generated content"}

        # FAQs
        brief["FAQs"] = {"output": "\n".join(faqs), "how": "List FAQs and AI-generated answers"}

        # Suggest other articles based on internal links
        brief["Consider writing topics on"] = {
            "output": "\n".join([l for l in internal_links if l not in serp_urls]),
            "how": "Recommend articles for the cluster based on semantic links"
        }

        progress_bar.progress(100)
        progress_text.text("Brief generated!")

        # ------------------------------
        # Display Table
        # ------------------------------
        st.subheader("SEO Content Brief Table")
        for section, data in brief.items():
            st.markdown(f"**{section}**")
            st.text_area(f"{section} output", value=data["output"], height=100, key=f"{section}_area")
            st.caption(f"How to fill: {data['how']}")

        # ------------------------------
        # Download DOCX
        # ------------------------------
        st.subheader("Download Brief")
        docx_file = create_docx(brief, keyword)
        st.download_button("Download DOCX", data=docx_file, file_name=f"{keyword}_brief.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
